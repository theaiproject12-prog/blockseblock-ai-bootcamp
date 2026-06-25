"""
Feature 4 starter: text extraction and chunking — YOUR IMPLEMENTATION GOES HERE.

The complete version lives in shared/ingestion.py (read it for reference).

ARCHITECTURE CONTEXT:
Four approaches exist for giving an LLM access to external documents.
This module builds the RAG path. The others are described here for context:

  RAG (what we're building): chunk → embed → vector DB → retrieve at query time.
      Scales to hundreds of documents. Requires retrieval to work correctly.
      Core limitation: "vibe retrieval" — similarity ≠ relevance.

  CAG (Cache-Augmented Generation): skip chunking — load ALL text into one
      big system prompt. No retrieval errors, but only practical for <50 pages.
      See the CAG PATTERN section at the bottom of this file.

  KAG (Knowledge-Augmented Generation): extract entities + relationships into
      a knowledge graph instead of a vector DB. Best for relational queries.
      Introduced in Feature 6.

  PageIndex (VectifyAI, github.com/VectifyAI/PageIndex, MIT licence):
      No chunking, no embeddings, no vector DB. Builds a hierarchical tree index
      (like a smart TOC) and uses LLM reasoning to navigate it. Achieved 98.7%
      on FinanceBench vs significantly lower for vector RAG. Best for long
      professional documents (financial reports, legal filings, technical manuals).
      See the PAGEINDEX NOTE section at the bottom of this file.

YOUR TASKS:
  Step 1: implement extract_text()   — three file-type branches (the critical one)
  Step 2: implement chunk_text()     — sentence accumulation + overlap (the critical one)
  Step 3 (read, don't write): study chunk_by_paragraph() and chunk_by_page()
          to understand how the other strategies differ

Everything else (chunk_by_paragraph reference, chunk_by_page reference,
CHUNKING_STRATEGIES dict, CAG pattern) is provided complete.
"""
import io
import re
from pathlib import Path
from typing import Callable


# =============================================================================
# Extraction
# =============================================================================

def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Extract plain text from the given file bytes.

    Dispatch on the file extension:

      .txt  — the bytes ARE the text; decode as UTF-8 (errors="replace").
               No extra library needed. One line:
               return file_bytes.decode("utf-8", errors="replace")

      .pdf  — use pypdf:
                from pypdf import PdfReader
                reader = PdfReader(io.BytesIO(file_bytes))
                return "\\n".join(page.extract_text() or "" for page in reader.pages)

      .docx — use python-docx:
                from docx import Document
                doc = Document(io.BytesIO(file_bytes))
                return "\\n".join(para.text for para in doc.paragraphs if para.text.strip())

      anything else — raise a clear ValueError:
                raise ValueError(f"Unsupported file type: '{ext}'. Supported: .txt, .pdf, .docx")

    Args:
      file_bytes: raw bytes of the uploaded file
      filename:   original filename (used only to determine extension)

    Returns:
      the extracted text as a single string
    """
    ext = Path(filename).suffix.lower()

    # TODO (Feature 4, Step 1a): Handle .txt files.
    # Hint: return file_bytes.decode("utf-8", errors="replace")
    if ext == ".txt":
        raise NotImplementedError("Implement .txt extraction — see the docstring above.")

    # TODO (Feature 4, Step 1b): Handle .pdf files.
    # Hint: from pypdf import PdfReader; use io.BytesIO(file_bytes)
    if ext == ".pdf":
        raise NotImplementedError("Implement .pdf extraction — see the docstring above.")

    # TODO (Feature 4, Step 1c): Handle .docx files.
    # Hint: from docx import Document; doc.paragraphs gives you a list of Paragraph objects
    if ext == ".docx":
        raise NotImplementedError("Implement .docx extraction — see the docstring above.")

    # TODO (Feature 4, Step 1d): Raise a clear error for unsupported types.
    raise ValueError(
        f"Unsupported file type: '{ext}'. Supported formats: .txt, .pdf, .docx."
    )


def extract_pages(file_bytes: bytes, filename: str) -> list[dict]:
    """
    Extract text per page, returning a list of {"page_number": int, "text": str} dicts.

    This function is provided complete — it builds on extract_text() and is
    used by chunk_by_page(). Read it to understand the page structure.

    For .pdf: one dict per page (page_number is 1-based).
    For .txt and .docx: single dict with page_number=1 and full text.
    """
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        return [
            {"page_number": i + 1, "text": page.extract_text() or ""}
            for i, page in enumerate(reader.pages)
        ]

    return [{"page_number": 1, "text": extract_text(file_bytes, filename)}]


# =============================================================================
# Chunking strategies
# =============================================================================

def chunk_text(text: str, chunk_size: int = 500, overlap: int = 50) -> list[str]:
    """
    Strategy: Sentence-aware fixed-size chunking.

    Split on sentence boundaries (. ! ?), group sentences until ~chunk_size
    characters, carry ~overlap characters from the tail of the previous chunk
    into the start of the next for context continuity.

    This is the DEFAULT strategy — works for most document types.

    Algorithm:
      1. (Given) Split text into sentences using regex.
      2. (Your work) For each sentence:
           a. If adding it would exceed chunk_size AND current is not empty:
              - emit " ".join(current) as a chunk
              - build overlap tail: walk reversed(current) until tail_len > overlap
              - set current = tail, current_len = tail_len
           b. Append the sentence: current.append(sentence); current_len += len(sentence) + 1
      3. (Given) Emit the final chunk.

    Args:
      text:       full document text
      chunk_size: target character length per chunk (default 500)
      overlap:    characters from previous chunk tail to carry forward (default 50)

    Returns:
      list of non-empty chunk strings; empty list if text is blank
    """
    # Step 1 (given — do not modify): split into sentences.
    sentences = re.split(r"(?<=[.!?])\s+", text.strip())
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return []

    chunks: list[str] = []
    current: list[str] = []   # sentences accumulated in the current chunk
    current_len = 0            # approximate character count of current chunk

    for sentence in sentences:
        slen = len(sentence)

        # TODO (Feature 4, Step 2a): if current is non-empty AND adding this
        # sentence would exceed chunk_size, emit the current chunk and reset.
        #
        # Emit:
        #   chunks.append(" ".join(current))
        #
        # Build overlap tail (carry last ~overlap chars into next chunk):
        #   tail, tail_len = [], 0
        #   for s in reversed(current):
        #       if tail_len + len(s) + 1 > overlap:
        #           break
        #       tail.insert(0, s)
        #       tail_len += len(s) + 1
        #   current, current_len = tail, tail_len

        if current and current_len + slen > chunk_size:
            raise NotImplementedError(
                "Implement chunk emission and overlap tail — see the TODO above (Step 2a)."
            )

        # TODO (Feature 4, Step 2b): append the sentence to current and update current_len.
        # current.append(sentence)
        # current_len += slen + 1   # +1 for the space join() inserts between sentences
        raise NotImplementedError(
            "Implement sentence accumulation — see the TODO above (Step 2b)."
        )

    # Step 3 (given — do not modify): emit the last chunk.
    if current:
        chunks.append(" ".join(current))

    return chunks


def chunk_by_paragraph(text: str, max_chunk_size: int = 800) -> list[str]:
    """
    Strategy: Paragraph-based chunking (reference implementation — read, don't rewrite).

    Split on double newlines (paragraph breaks). If a paragraph exceeds
    max_chunk_size, fall back to sentence-aware splitting for that paragraph only.

    Best for: formal structured documents — policies, manuals, legal text.

    This is provided so you can compare it with chunk_text() above:
    - chunk_text() splits on SENTENCES, groups until size limit
    - chunk_by_paragraph() splits on PARAGRAPHS, respects natural breaks
    Same idea, different boundary signal.
    """
    paragraphs = [p.strip() for p in re.split(r"\n\n+", text) if p.strip()]
    chunks: list[str] = []

    for para in paragraphs:
        if len(para) <= max_chunk_size:
            chunks.append(para)
        else:
            # TODO (Feature 4, optional stretch): call chunk_text(para, chunk_size=max_chunk_size)
            # and extend chunks with the result. One line.
            # This is already implemented in shared/ingestion.py — read it for reference.
            chunks.extend(chunk_text(para, chunk_size=max_chunk_size))

    return chunks


def chunk_by_page(pages: list[dict], max_page_size: int = 2000) -> list[dict]:
    """
    Strategy: pageIndex — one chunk per page (reference implementation — read, don't rewrite).

    Each returned dict has "text", "page_number", and "chunk_index". If a page
    exceeds max_page_size, it is sub-chunked with chunk_text() while preserving
    the page_number.

    When to use: financial reports, legal filings, academic papers — anywhere
    "see page X" is a meaningful citation. The page_number in metadata lets
    the assistant tell the user exactly which page an answer came from.

    Study this implementation — it shows how metadata flows from extraction
    through chunking into the stored Chunk objects and eventually into the UI.
    """
    result: list[dict] = []
    chunk_index = 0

    for page in pages:
        text = page["text"].strip()
        page_num = page["page_number"]

        if not text:
            continue

        if len(text) <= max_page_size:
            result.append({"text": text, "page_number": page_num, "chunk_index": chunk_index})
            chunk_index += 1
        else:
            for sub in chunk_text(text, chunk_size=max_page_size):
                result.append({"text": sub, "page_number": page_num, "chunk_index": chunk_index})
                chunk_index += 1

    return result


# =============================================================================
# CAG ALTERNATIVE (Cache-Augmented Generation) — documented pattern, not a TODO
#
# Instead of chunking, concatenate ALL document text and pass it as part of
# the system prompt. No vector DB, no retrieval step, no chunking decisions.
# The model reads everything and answers from full context.
#
# When to use:
#   • Total document set < ~50 pages (fits in a large context window)
#   • Documents change infrequently (cache must be recomputed on each change)
#   • Zero retrieval errors are required (you can't afford a missed chunk)
#   • You're using GPT-4o (128K tokens) or Gemini 1.5 Pro (1M tokens)
#
# Pattern (not wired up — read and understand, don't implement):
#
#   all_text = "\n\n---\n\n".join(
#       extract_text(file_bytes, filename) for file_bytes, filename in documents
#   )
#   system_prompt = (
#       "Here is all relevant knowledge:\n\n"
#       f"{all_text}\n\n"
#       "Answer the user's question using only the knowledge above."
#   )
#   result = await call_llm([
#       {"role": "system", "content": system_prompt},
#       {"role": "user",   "content": user_query},
#   ])
#
# Compare with RAG (what we're building):
#   CAG sends everything; retrieval errors = 0. Cost grows with doc count.
#   RAG retrieves 2-3 chunks; cost stays constant. Retrieval can miss things.
#   At < 20 pages, try CAG. At 50+ pages, RAG scales better.
# =============================================================================


# =============================================================================
# PAGEINDEX NOTE (Vectorless, Reasoning-based RAG) — reference, not implemented here.
#
# VectifyAI's PageIndex (github.com/VectifyAI/PageIndex, MIT) takes a different
# approach from ALL chunking strategies in this file:
#
#   1. Build a HIERARCHICAL TREE INDEX — like a smart table of contents.
#      Each tree node has: title, page range, LLM-generated section summary.
#   2. Retrieve by REASONING over the tree — an LLM agent navigates the tree
#      the way a human expert flips through a complex document.
#      No vector similarity. No embeddings. No chunk boundaries.
#
# "similarity ≠ relevance" — the key insight. Vector RAG returns semantically
# similar passages; PageIndex returns the passage an expert would find.
# Achieved 98.7% accuracy on FinanceBench vs significantly lower for vector RAG.
#
# When to use instead of this module:
#   • Long professional docs: financial reports, legal filings, technical manuals
#   • Multi-step reasoning needed to find the right section
#   • Vector RAG keeps returning wrong answers despite tuning
#
# Integration path for Feature 6 Smart Router:
#   pip install -r requirements.txt  # from github.com/VectifyAI/PageIndex
#   python run_pageindex.py --pdf_path your_doc.pdf  # builds tree JSON
#   In Feature 6: route professional-doc queries → PageIndex tree search
#                 route general queries          → vector RAG (Features 4-6)
#   Cloud service + MCP server at pageindex.ai
# =============================================================================


# =============================================================================
# Strategy registry — fully wired, no TODOs here.
# The upload endpoint uses this dict to select a strategy by name.
# Only "sentence" and "paragraph" are exposed in the UI — chunk_by_page()
# exists above as reference code but is not included here.
# =============================================================================

def _sentence_strategy(text: str, pages: list[dict]) -> list[dict]:
    return [
        {"text": c, "page_number": None, "chunk_index": i}
        for i, c in enumerate(chunk_text(text))
    ]


def _paragraph_strategy(text: str, pages: list[dict]) -> list[dict]:
    return [
        {"text": c, "page_number": None, "chunk_index": i}
        for i, c in enumerate(chunk_by_paragraph(text))
    ]


CHUNKING_STRATEGIES: dict[str, Callable[[str, list[dict]], list[dict]]] = {
    "sentence":  _sentence_strategy,
    "paragraph": _paragraph_strategy,
    # "page": chunk_by_page() above is the per-page alternative.
    # For a more powerful page-based approach see PAGEINDEX NOTE above.
}
